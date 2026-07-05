"""
Génération DOCX des CV.

Le rendu final est piloté par les données du CV:
- identité et contact,
- sections activées/désactivées,
- ordre des sections,
- style visuel du template choisi.
"""
import io
import re
import subprocess
import tempfile
from pathlib import Path
from zipfile import ZipFile

from django.conf import settings
from django.utils import timezone

from django.core.files.base import ContentFile
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from cvs.system_bins import candidate_libreoffice_binaries, libreoffice_not_found_message


SECTION_ORDER = ["profile", "experiences", "education", "skills", "languages", "hobbies"]
SIDEBAR_SECTIONS = {"skills", "languages", "hobbies"}

SECTION_LABELS = {
    "profile": "Profil",
    "experiences": "Expériences professionnelles",
    "education": "Formation",
    "skills": "Compétences",
    "languages": "Langues",
    "hobbies": "Loisirs",
}

TEMPLATE_STYLES = {
    "modele-simple": {"accent": "1F2937", "variant": "simple", "font": "Calibri", "compact": False},
    "modele-classique": {"accent": "1D4ED8", "variant": "classic", "font": "Calibri", "compact": False},
    "modele-moderne": {"accent": "0F766E", "variant": "modern", "font": "Aptos", "compact": False},
    "modele-compact": {"accent": "7C2D12", "variant": "compact", "font": "Arial", "compact": True},
    "modele-executif": {"accent": "111827", "variant": "executive", "font": "Georgia", "compact": False},
    "modele-creatif": {"accent": "9F1239", "variant": "creative", "font": "Aptos", "compact": False},
    "classique-elegant": {"accent": "1D4ED8", "variant": "classic", "font": "Calibri", "compact": False},
    "moderne-dynamique": {"accent": "0F766E", "variant": "modern", "font": "Aptos", "compact": False},
    "minimaliste-epure": {"accent": "111827", "variant": "executive", "font": "Georgia", "compact": False},
    "creatif-moderne": {"accent": "9F1239", "variant": "creative", "font": "Aptos", "compact": False},
    "colore-vibrant": {"accent": "9F1239", "variant": "creative", "font": "Aptos", "compact": False},
}

CATEGORY_STYLES = {
    "classic": TEMPLATE_STYLES["modele-classique"],
    "modern": TEMPLATE_STYLES["modele-moderne"],
    "creative": TEMPLATE_STYLES["modele-creatif"],
    "minimal": TEMPLATE_STYLES["modele-executif"],
    "colorful": TEMPLATE_STYLES["colore-vibrant"],
}


GALLERY_VARIANTS = {
    0: "sidebar",
    1: "header",
    2: "minimal",
    3: "rail",
    4: "executive",
    5: "sidebar",
    6: "header",
    7: "minimal",
    8: "rail",
    9: "executive",
}


def _gallery_style(template):
    match = re.match(r"^galerie-cv-(\d+)$", template.slug or "")
    if not match:
        return None
    index = int(match.group(1)) - 1
    base = CATEGORY_STYLES.get(getattr(template, "category", None), TEMPLATE_STYLES["modele-simple"]).copy()
    base["variant"] = GALLERY_VARIANTS[index % 10]
    base["compact"] = base["variant"] == "minimal"
    return base


def _style_for(cv):
    gallery_style = _gallery_style(cv.template)
    if gallery_style:
        return gallery_style
    return (
        TEMPLATE_STYLES.get(cv.template.slug)
        or CATEGORY_STYLES.get(getattr(cv.template, "category", None))
        or TEMPLATE_STYLES["modele-simple"]
    )


def _rgb(hex_color):
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _safe_text(value):
    return str(value or "").strip()


def _clean_items(items):
    cleaned = []
    for item in items or []:
        if isinstance(item, dict):
            if any(_safe_text(value) for value in item.values() if not isinstance(value, list)):
                cleaned.append(item)
        elif _safe_text(item):
            cleaned.append(_safe_text(item))
    return cleaned


def _photo_path(data):
    value = _safe_text(data.get("photo_url") or data.get("photo"))
    if not value:
        return None
    marker = settings.MEDIA_URL or "/media/"
    if marker and marker in value:
        relative = value.split(marker, 1)[1]
        path = Path(settings.MEDIA_ROOT) / relative
        return path if path.exists() else None
    path = Path(value)
    return path if path.exists() else None


def _add_photo(container, data, width=0.95, align=None):
    path = _photo_path(data)
    if not path:
        return False
    paragraph = container.add_paragraph()
    if align:
        paragraph.alignment = align
    run = paragraph.add_run()
    try:
        run.add_picture(str(path), width=Inches(width))
    except Exception:
        return False
    paragraph.paragraph_format.space_after = Pt(6)
    return True


def _enabled_sections(data):
    enabled = {section: True for section in SECTION_ORDER}
    if isinstance(data.get("enabled_sections"), dict):
        enabled.update(data["enabled_sections"])
    return enabled


def _ordered_sections(data):
    saved = data.get("section_order") if isinstance(data.get("section_order"), list) else []
    known = [section for section in saved if section in SECTION_ORDER]
    return known + [section for section in SECTION_ORDER if section not in known]


def _sections_to_render(data, include_sidebar=True):
    enabled = _enabled_sections(data)
    sections = [section for section in _ordered_sections(data) if enabled.get(section)]
    if include_sidebar:
        return sections
    return [section for section in sections if section not in SIDEBAR_SECTIONS]


def _set_run(run, style, size=10, bold=False, color=None, italic=False):
    run.font.name = style["font"]
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = _rgb(color)


def _add_paragraph(container, text="", style=None, size=10, bold=False, color=None, italic=False, align=None, space_after=4):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(_safe_text(text))
    _set_run(run, style, size=size, bold=bold, color=color, italic=italic)
    return paragraph


def _add_heading(container, label, style):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8 if not style["compact"] else 4)
    paragraph.paragraph_format.space_after = Pt(5 if not style["compact"] else 2)
    run = paragraph.add_run(label.upper())
    _set_run(run, style, size=10 if not style["compact"] else 8.5, bold=True, color=style["accent"])
    return paragraph


def _add_bullet(container, text, style, size=9.5):
    paragraph = container.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(_safe_text(text))
    _set_run(run, style, size=size)
    return paragraph


def _contact_parts(data):
    return [
        _safe_text(data.get("phone")),
        _safe_text(data.get("email")),
        _safe_text(data.get("address")),
        _safe_text(data.get("linkedin")),
        _safe_text(data.get("github")),
        _safe_text(data.get("portfolio")),
        _safe_text(data.get("driving_license")),
    ]


def _add_header(container, data, style):
    first_name = _safe_text(data.get("first_name"))
    last_name = _safe_text(data.get("last_name"))
    full_name = " ".join([part for part in [first_name, last_name] if part]) or "Mon CV"
    align = WD_ALIGN_PARAGRAPH.CENTER if style["variant"] == "executive" else WD_ALIGN_PARAGRAPH.LEFT

    if style["variant"] in {"executive", "header", "minimal"}:
        _add_photo(container, data, width=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)

    _add_paragraph(container, full_name, style, size=24 if not style["compact"] else 18, bold=True, color=style["accent"], align=align, space_after=2)
    if data.get("job_title"):
        _add_paragraph(container, data.get("job_title"), style, size=12 if not style["compact"] else 10, bold=True, align=align, space_after=3)

    contact = " | ".join([part for part in _contact_parts(data) if part])
    if contact:
        _add_paragraph(container, contact, style, size=8.5 if not style["compact"] else 7.5, color="4B5563", align=align, space_after=10)


def _add_profile(container, data, style):
    profile = _safe_text(data.get("profile"))
    if not profile:
        return
    _add_heading(container, SECTION_LABELS["profile"], style)
    _add_paragraph(container, profile, style, size=9.5 if not style["compact"] else 8.2, space_after=5)


def _add_experiences(container, data, style):
    experiences = _clean_items(data.get("experiences"))
    if not experiences:
        return
    _add_heading(container, SECTION_LABELS["experiences"], style)
    for exp in experiences:
        title = " | ".join([part for part in [_safe_text(exp.get("job_title")), _safe_text(exp.get("company"))] if part])
        meta = " | ".join([part for part in [_safe_text(exp.get("period")), _safe_text(exp.get("location")), _safe_text(exp.get("type"))] if part])
        if title:
            _add_paragraph(container, title, style, size=10 if not style["compact"] else 8.6, bold=True, space_after=1)
        if meta:
            _add_paragraph(container, meta, style, size=8.3 if not style["compact"] else 7.4, color="6B7280", space_after=1)
        for mission in exp.get("missions") or []:
            if _safe_text(mission):
                _add_bullet(container, mission, style, size=8.8 if not style["compact"] else 7.6)


def _add_education(container, data, style):
    education = _clean_items(data.get("education"))
    if not education:
        return
    _add_heading(container, SECTION_LABELS["education"], style)
    for edu in education:
        title = _safe_text(edu.get("degree"))
        meta = " | ".join([part for part in [_safe_text(edu.get("institution")), _safe_text(edu.get("location")), _safe_text(edu.get("period"))] if part])
        if title:
            _add_paragraph(container, title, style, size=9.5 if not style["compact"] else 8.2, bold=True, space_after=1)
        if meta:
            _add_paragraph(container, meta, style, size=8.3 if not style["compact"] else 7.4, color="6B7280", space_after=2)


def _add_skills(container, data, style):
    skills = _clean_items(data.get("skills"))
    if not skills:
        return
    _add_heading(container, SECTION_LABELS["skills"], style)
    _add_paragraph(container, " · ".join(skills), style, size=9 if not style["compact"] else 7.8, space_after=4)


def _add_languages(container, data, style):
    languages = _clean_items(data.get("languages"))
    if not languages:
        return
    _add_heading(container, SECTION_LABELS["languages"], style)
    for language in languages:
        label = " - ".join([part for part in [_safe_text(language.get("language")), _safe_text(language.get("level"))] if part])
        if label:
            _add_bullet(container, label, style, size=8.8 if not style["compact"] else 7.6)


def _add_hobbies(container, data, style):
    hobbies = _clean_items(data.get("hobbies"))
    if not hobbies:
        return
    _add_heading(container, SECTION_LABELS["hobbies"], style)
    _add_paragraph(container, ", ".join(hobbies), style, size=9 if not style["compact"] else 7.8, space_after=4)


def _add_extra_sections(container, data, style):
    for section in _clean_items(data.get("extra_sections")):
        title = _safe_text(section.get("title"))
        items = _clean_items(section.get("items"))
        if not title or not items:
            continue
        _add_heading(container, title, style)
        for item in items:
            _add_bullet(container, item, style, size=8.8 if not style["compact"] else 7.6)


SECTION_RENDERERS = {
    "profile": _add_profile,
    "experiences": _add_experiences,
    "education": _add_education,
    "skills": _add_skills,
    "languages": _add_languages,
    "hobbies": _add_hobbies,
}


def _render_sections(container, data, style, sections):
    for section in sections:
        SECTION_RENDERERS[section](container, data, style)


def _configure_document(doc, style):
    section = doc.sections[0]
    section.top_margin = Inches(0.55 if style["compact"] else 0.7)
    section.bottom_margin = Inches(0.55 if style["compact"] else 0.7)
    section.left_margin = Inches(0.6 if style["compact"] else 0.75)
    section.right_margin = Inches(0.6 if style["compact"] else 0.75)

    normal = doc.styles["Normal"]
    normal.font.name = style["font"]
    normal.font.size = Pt(9 if style["compact"] else 10)


def _build_single_column_doc(data, style):
    doc = Document()
    _configure_document(doc, style)
    _add_header(doc, data, style)
    _render_sections(doc, data, style, _sections_to_render(data))
    _add_extra_sections(doc, data, style)
    return doc


def _build_two_column_doc(data, style):
    doc = Document()
    _configure_document(doc, style)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    left.width = Inches(2.15)
    right.width = Inches(4.35)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    _add_header(right, data, style)
    _render_sections(right, data, style, _sections_to_render(data, include_sidebar=False))
    _add_extra_sections(right, data, style)

    _add_photo(left, data, width=1.18, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_heading(left, "Contact", style)
    for part in _contact_parts(data):
        if part:
            _add_paragraph(left, part, style, size=8.3, space_after=3)
    sidebar = [section for section in ["skills", "languages", "hobbies"] if _enabled_sections(data).get(section)]
    _render_sections(left, data, style, sidebar)
    return doc


def _build_document(cv):
    style = _style_for(cv)
    data = cv.data or {}
    if style["variant"] in {"modern", "creative", "sidebar", "rail"}:
        return _build_two_column_doc(data, style)
    return _build_single_column_doc(data, style)



def _template_path(cv):
    name = (cv.template.docx_filename or "").strip()
    if not name:
        return None
    relative = Path(name)
    if relative.is_absolute() or ".." in relative.parts:
        return None
    path = settings.BASE_DIR / "cvs" / "docx_templates" / relative
    return path if path.exists() else None


def _template_has_placeholders(path):
    try:
        with ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        return "{{" in xml or "{%" in xml
    except Exception:
        return False


def _formatted_context(data):
    data = data or {}
    experiences = _clean_items(data.get("experiences"))
    education = _clean_items(data.get("education"))
    languages = _clean_items(data.get("languages"))
    skills = _clean_items(data.get("skills"))
    hobbies = _clean_items(data.get("hobbies"))

    def exp_text(exp):
        lines = []
        title = " - ".join([part for part in [_safe_text(exp.get("job_title")), _safe_text(exp.get("company"))] if part])
        meta = " | ".join([part for part in [_safe_text(exp.get("period")), _safe_text(exp.get("location")), _safe_text(exp.get("type"))] if part])
        if title:
            lines.append(title)
        if meta:
            lines.append(meta)
        for mission in exp.get("missions") or []:
            if _safe_text(mission):
                lines.append(f"• {_safe_text(mission)}")
        return "\n".join(lines)

    def edu_text(edu):
        title = _safe_text(edu.get("degree"))
        meta = " | ".join([part for part in [_safe_text(edu.get("institution")), _safe_text(edu.get("location")), _safe_text(edu.get("period"))] if part])
        return "\n".join([part for part in [title, meta] if part])

    context = {**data}
    context.update({
        "full_name": " ".join([part for part in [_safe_text(data.get("first_name")), _safe_text(data.get("last_name"))] if part]),
        "experiences": experiences,
        "education": education,
        "skills": skills,
        "languages": languages,
        "hobbies": hobbies,
        "experiences_text": "\n\n".join(exp_text(exp) for exp in experiences),
        "education_text": "\n\n".join(edu_text(edu) for edu in education),
        "skills_text": " • ".join(skills),
        "languages_text": "\n".join(" - ".join([part for part in [_safe_text(item.get("language")), _safe_text(item.get("level"))] if part]) for item in languages),
        "hobbies_text": ", ".join(hobbies),
    })
    return context


def _build_docx_template_document(cv):
    path = _template_path(cv)
    if not path or not _template_has_placeholders(path):
        return None
    try:
        from docxtpl import DocxTemplate
    except Exception:
        return None
    template = DocxTemplate(path)
    template.render(_formatted_context(cv.data or {}))
    return template


def _save_document_to_bytes(document):
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _convert_docx_to_pdf(docx_bytes, base_name):
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        docx_path = tmp / f"{base_name}.docx"
        outdir = tmp / "out"
        profile = tmp / "profile"
        outdir.mkdir()
        profile.mkdir()
        docx_path.write_bytes(docx_bytes)
        pdf_path = outdir / f"{base_name}.pdf"
        binaries = candidate_libreoffice_binaries()
        if not binaries:
            raise FileNotFoundError(libreoffice_not_found_message())

        errors = []
        for binary in binaries:
            pdf_path.unlink(missing_ok=True)
            command = [
                binary,
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(outdir),
                str(docx_path),
            ]
            completed = subprocess.run(command, capture_output=True, text=True, timeout=60, check=False)
            if completed.returncode == 0 and pdf_path.exists():
                return pdf_path.read_bytes()
            errors.append((completed.stderr or completed.stdout or f"{binary}: conversion PDF impossible").strip())
        raise RuntimeError(" | ".join(errors))

def _filename_base(cv):
    base = re.sub(r"[^a-zA-Z0-9]+", "_", cv.title or "cv").strip("_").lower() or "cv"
    return f"{base}_{cv.id}"


def _filename(cv):
    return f"{_filename_base(cv)}.docx"


def _pdf_filename(cv):
    return f"{_filename_base(cv)}.pdf"


def _render_html_pdf_bytes(cv):
    try:
        from cvs.renderers.html import HTMLRendererUnavailable, render_html_cv_pdf_bytes, supports_html_renderer
    except Exception:
        return None
    if not supports_html_renderer(cv.template):
        return None
    try:
        return render_html_cv_pdf_bytes(cv, base_name=_filename_base(cv))
    except HTMLRendererUnavailable:
        return None


def _gallery_model_index(template):
    match = re.match(r"^galerie-cv-(\d+)$", template.slug or "")
    return int(match.group(1)) - 1 if match else None


def _gallery_photo_path(data):
    return _photo_path(data) or settings.BASE_DIR.parent / "frontend" / "src" / "assets" / "cv-photos" / "normalized" / "cv-photo-01.png"


def _gallery_cv_data(data):
    data = data or {}
    full_name = " ".join([part for part in [_safe_text(data.get("first_name")), _safe_text(data.get("last_name"))] if part]) or "Mon CV"
    contact = " | ".join([part for part in _contact_parts(data) if part])
    experiences = []
    for exp in _clean_items(data.get("experiences"))[:3]:
        role = _safe_text(exp.get("job_title")) or "Expérience professionnelle"
        company = _safe_text(exp.get("company"))
        meta = " | ".join([part for part in [company, _safe_text(exp.get("period"))] if part])
        missions = _clean_items(exp.get("missions"))
        detail = missions[0] if missions else _safe_text(exp.get("type"))
        experiences.append((role, meta, detail))
    if not experiences:
        experiences = [("Expérience professionnelle", "", "Ajoutez vos missions et résultats principaux.")]

    education = []
    for edu in _clean_items(data.get("education"))[:2]:
        title = _safe_text(edu.get("degree")) or "Formation"
        meta = " | ".join([part for part in [_safe_text(edu.get("institution")), _safe_text(edu.get("period"))] if part])
        education.append((title, meta))
    if not education:
        education = [("Formation", "Etablissement | Période")]

    languages = [
        " - ".join([part for part in [_safe_text(item.get("language")), _safe_text(item.get("level"))] if part])
        for item in _clean_items(data.get("languages"))
    ]
    extra = {str(section.get("title", "")).lower(): _clean_items(section.get("items")) for section in _clean_items(data.get("extra_sections"))}

    def section_items(*tokens, fallback=None):
        for title, items in extra.items():
            if any(token in title for token in tokens) and items:
                return items
        return fallback or []

    return {
        "name": full_name,
        "role": _safe_text(data.get("job_title")) or "Titre professionnel",
        "contact": contact or "Téléphone | Email | Ville",
        "profile": _safe_text(data.get("profile")) or "Profil professionnel synthétique, orienté résultats et adapté au poste ciblé.",
        "experiences": experiences,
        "education": education,
        "skills": _clean_items(data.get("skills"))[:8] or ["Compétence clé", "Organisation", "Communication"],
        "languages": [item for item in languages if item] or ["Français - Courant"],
        "tools": section_items("outil", "logiciel", fallback=_clean_items(data.get("skills"))[3:8] or ["Pack Office"]),
        "achievements": section_items("réalisation", "realisation", "projet", fallback=["Résultat ou projet important à valoriser."]),
        "certifications": section_items("certification", "certificat", fallback=["Certification ou formation complémentaire."]),
        "references": section_items("référence", "reference", fallback=["Références disponibles sur demande."]),
        "interests": _clean_items(data.get("hobbies")) or section_items("loisir", "intérêt", "interet", fallback=["Veille métier"]),
    }


def _render_gallery_pdf_bytes(cv):
    index = _gallery_model_index(cv.template)
    if index is None:
        return None
    try:
        from templates.management.commands.generate_preview_gallery import A4_SIZE, LAYOUTS, PALETTES
        from PIL import Image, ImageDraw
    except Exception:
        return None

    palette = PALETTES[index % len(PALETTES)]
    _, layout = LAYOUTS[index % len(LAYOUTS)]
    canvas = Image.new("RGB", A4_SIZE, "#ffffff")
    draw = ImageDraw.Draw(canvas)
    layout(canvas, draw, _gallery_cv_data(cv.data or {}), _gallery_photo_path(cv.data or {}), palette, index % len(LAYOUTS))
    buffer = io.BytesIO()
    canvas.save(buffer, "PDF", resolution=150)
    buffer.seek(0)
    return buffer.getvalue()


def _build_output_document(cv):
    templated = _build_docx_template_document(cv)
    if templated is not None:
        return templated
    return _build_document(cv)


def generate_cv_documents(cv):
    """Génère le DOCX puis le PDF final d'un CV."""
    template = cv.template
    document = _build_output_document(cv)
    docx_bytes = _save_document_to_bytes(document)
    # Source unique : le PDF est TOUJOURS le rendu HTML/CSS (identique à l'aperçu).
    # Le DOCX n'est conservé que pour le téléchargement Word, jamais comme source du PDF.
    pdf_bytes = _render_html_pdf_bytes(cv) or _convert_docx_to_pdf(docx_bytes, _filename_base(cv))

    cv.generated_file.save(_filename(cv), ContentFile(docx_bytes), save=False)
    cv.generated_pdf.save(_pdf_filename(cv), ContentFile(pdf_bytes), save=False)
    cv.status = cv.STATUS_GENERATED
    cv.generated_at = timezone.now()
    cv.save(update_fields=["generated_file", "generated_pdf", "status", "generated_at", "updated_at"])
    cv._page_count = _count_pdf_pages(pdf_bytes)
    return cv.generated_pdf


def _count_pdf_pages(pdf_bytes):
    try:
        from pypdf import PdfReader

        return len(PdfReader(io.BytesIO(pdf_bytes)).pages)
    except Exception:
        return 1


def generate_cv_docx(cv):
    """Compatibilité avec l'ancien endpoint: génère maintenant DOCX + PDF."""
    generate_cv_documents(cv)
    return cv.generated_file
