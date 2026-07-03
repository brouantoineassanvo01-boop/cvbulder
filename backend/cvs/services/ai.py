import base64
import io
import json
import mimetypes
import re
import shutil
import subprocess
import tempfile
import time
import urllib.error
import urllib.request
from difflib import SequenceMatcher
from html.parser import HTMLParser
from pathlib import Path
from zipfile import ZipFile

from django.conf import settings
from PIL import Image, ImageEnhance, ImageFilter, ImageOps, UnidentifiedImageError
from templates.services.design_contracts import template_design_contract


class AIServiceError(RuntimeError):
    def __init__(self, message, code="ai_service_error", status_code=500, raw_detail=""):
        super().__init__(message)
        self.code = code
        self.status_code = status_code
        self.raw_detail = raw_detail


class _HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []
        self.skip = False

    def handle_starttag(self, tag, attrs):
        if tag in {"script", "style", "noscript"}:
            self.skip = True

    def handle_endtag(self, tag):
        if tag in {"script", "style", "noscript"}:
            self.skip = False

    def handle_data(self, data):
        if not self.skip:
            text = data.strip()
            if text:
                self.parts.append(text)

    def text(self):
        return re.sub(r"\s+", " ", " ".join(self.parts)).strip()


def _shorten(text, limit=12000):
    lines = []
    for line in str(text or "").replace("\r", "\n").split("\n"):
        cleaned = re.sub(r"[ \t]+", " ", line).strip()
        if cleaned:
            lines.append(cleaned)
    text = "\n".join(lines).strip()
    return text[:limit]


def _safe_text(value):
    return str(value or "").strip()


def _clean_items(items):
    cleaned = []
    for item in items or []:
        if isinstance(item, dict):
            scalar_values = [value for value in item.values() if not isinstance(value, list)]
            list_values = [value for value in item.values() if isinstance(value, list)]
            if any(_safe_text(value) for value in scalar_values) or any(_clean_items(value) for value in list_values):
                cleaned.append(item)
        elif _safe_text(item):
            cleaned.append(_safe_text(item))
    return cleaned


def _extract_docx(path):
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(paragraphs)


def _text_quality(text):
    text = str(text or "")
    words = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9][A-Za-zÀ-ÖØ-öø-ÿ0-9'._+-]{1,}", text)
    alnum_count = sum(1 for char in text if char.isalnum())
    line_count = len([line for line in text.splitlines() if line.strip()])
    has_contact = bool(re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, re.I)) or bool(
        re.search(r"(?:\+\d{1,4}[\s.-]?)?(?:\d[\s.-]?){8,}", text)
    )
    weird_count = sum(1 for char in text if not (char.isalnum() or char.isspace() or char in ".,;:!?@/\\-+()[]'\"&%#"))
    weird_ratio = weird_count / max(len(text), 1)
    score = min(100, int(len(words) * 1.8 + alnum_count / 18 + line_count * 2 + (18 if has_contact else 0)))
    if weird_ratio > 0.12:
        score = max(0, score - int(weird_ratio * 100))
    usable = (len(words) >= 28 and alnum_count >= 140 and weird_ratio <= 0.2) or (has_contact and len(words) >= 14)
    return {
        "usable": usable,
        "score": score,
        "word_count": len(words),
        "alnum_count": alnum_count,
        "line_count": line_count,
        "has_contact": has_contact,
        "weird_ratio": weird_ratio,
    }


def _usable_text(text):
    return _text_quality(text)["usable"]


def _has_contact_hints(text):
    hints = _source_text_hints(text)
    return any(
        hints.get(key)
        for key in ["possible_email", "possible_phone", "possible_linkedin", "possible_github", "possible_portfolio"]
    )


def _merge_extracted_text(primary, secondary):
    primary = str(primary or "").strip()
    secondary = str(secondary or "").strip()
    if not secondary:
        return primary
    if not primary:
        return secondary
    if secondary in primary:
        return primary
    if primary in secondary:
        return secondary
    return f"{primary}\n\n--- OCR complementaire ---\n{secondary}"


def _extract_pdf_contact_ocr(path):
    pdftoppm = shutil.which("pdftoppm")
    tesseract = shutil.which("tesseract")
    if not pdftoppm or not tesseract:
        return ""

    languages = getattr(settings, "CV_OCR_LANGUAGES", "fra+eng") or "fra+eng"
    dpi = min(150, max(120, int(getattr(settings, "CV_OCR_DPI", 180))))
    texts = []
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp = Path(tmpdir)
            prefix = tmp / "contact-page"
            render = subprocess.run(
                [pdftoppm, "-png", "-singlefile", "-r", str(dpi), "-f", "1", "-l", "1", str(path), str(prefix)],
                capture_output=True,
                text=True,
                timeout=20,
                check=False,
            )
            image_path = prefix.with_suffix(".png")
            if render.returncode != 0 or not image_path.exists():
                return ""

            image = Image.open(image_path).convert("RGB")
            regions = [("full", image)]
            for name, region in regions:
                region_path = tmp / f"{name}.png"
                region.save(region_path, "PNG", optimize=True)
                text = _run_tesseract(region_path, languages, 6, timeout=14)
                if text.strip():
                    texts.append(text)
    except Exception:
        return ""
    return "\n".join(texts)


def _extract_pdf(path):
    best_text = ""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "source.txt"
            completed = subprocess.run(
                ["pdftotext", "-layout", "-f", "1", "-l", "6", str(path), str(output_path)],
                capture_output=True,
                text=True,
                timeout=20,
                check=False,
            )
            if completed.returncode == 0 and output_path.exists():
                text = output_path.read_text(encoding="utf-8", errors="ignore")
                if _usable_text(text):
                    if _has_contact_hints(text):
                        return text
                    ocr_text = _extract_pdf_contact_ocr(path)
                    return _merge_extracted_text(text, ocr_text)
                best_text = text if len(text) > len(best_text) else best_text
    except Exception:
        pass
    try:
        from pypdf import PdfReader
    except Exception:
        ocr_text = _extract_pdf_ocr(path)
        return ocr_text or best_text
    try:
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages[:6]]
        text = "\n".join(pages)
        if _usable_text(text):
            if _has_contact_hints(text):
                return text
            ocr_text = _extract_pdf_contact_ocr(path)
            return _merge_extracted_text(text, ocr_text)
        best_text = text if len(text) > len(best_text) else best_text
    except Exception:
        pass
    ocr_text = _extract_pdf_ocr(path)
    return ocr_text or best_text


def _extract_pdf_ocr(path):
    tesseract = shutil.which("tesseract")
    pdftoppm = shutil.which("pdftoppm")
    if not tesseract or not pdftoppm:
        return ""

    max_pages = max(1, int(getattr(settings, "CV_OCR_MAX_PAGES", 6)))
    dpi = max(96, int(getattr(settings, "CV_OCR_DPI", 180)))
    languages = getattr(settings, "CV_OCR_LANGUAGES", "fra+eng") or "fra+eng"
    texts = []
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            prefix = Path(tmpdir) / "ocr-page"
            render = subprocess.run(
                [pdftoppm, "-png", "-r", str(dpi), "-f", "1", "-l", str(max_pages), str(path), str(prefix)],
                capture_output=True,
                text=True,
                timeout=45,
                check=False,
            )
            if render.returncode != 0:
                return ""
            for image_path in sorted(Path(tmpdir).glob("ocr-page-*.png")):
                raw_text = _run_tesseract(image_path, languages, 6, timeout=20)
                if _usable_text(raw_text):
                    texts.append(raw_text)
                    continue
                prepared_path = image_path.with_name(f"{image_path.stem}-prepared.png")
                _preprocess_ocr_image(image_path, prepared_path)
                prepared_text = _best_ocr_text_for_image(prepared_path, languages)
                text = _merge_extracted_text(prepared_text, raw_text)
                if text.strip():
                    texts.append(text)
    except Exception:
        return ""
    return "\n".join(texts)


def _preprocess_ocr_image(source_path, output_path):
    image = Image.open(source_path)
    image = ImageOps.exif_transpose(image.convert("L"))
    width, height = image.size
    if min(width, height) < 1200:
        scale = min(2.0, 1200 / max(min(width, height), 1))
        image = image.resize((int(width * scale), int(height * scale)), Image.Resampling.LANCZOS)
    image = ImageOps.autocontrast(image)
    image = image.filter(ImageFilter.MedianFilter(size=3))
    image = ImageEnhance.Contrast(image).enhance(1.45)
    image = ImageEnhance.Sharpness(image).enhance(1.35)
    image.save(output_path, "PNG", optimize=True)
    return output_path


def _run_tesseract(image_path, languages, psm, timeout=45):
    completed = subprocess.run(
        ["tesseract", str(image_path), "stdout", "-l", languages, "--psm", str(psm)],
        capture_output=True,
        text=True,
        timeout=timeout,
        check=False,
    )
    if completed.returncode != 0:
        return ""
    return completed.stdout


def _rotated_image(source_path, output_path, degrees):
    image = Image.open(source_path)
    image = image.rotate(degrees, expand=True, fillcolor=255)
    image.save(output_path, "PNG", optimize=True)
    return output_path


def _best_ocr_text_for_image(image_path, languages):
    candidates = []
    for psm in (1, 4, 6):
        text = _run_tesseract(image_path, languages, psm)
        if text.strip():
            candidates.append(text)
    try:
        inverted_path = image_path.with_name(f"{image_path.stem}-inverted.png")
        inverted = ImageOps.invert(Image.open(image_path).convert("L"))
        inverted = ImageOps.autocontrast(inverted)
        inverted.save(inverted_path, "PNG", optimize=True)
        for psm in (4, 6):
            text = _run_tesseract(inverted_path, languages, psm)
            if text.strip():
                candidates.append(text)
    except Exception:
        pass
    best = max(candidates, key=lambda item: _text_quality(item)["score"], default="")
    if _usable_text(best):
        return best

    rotated_candidates = [best] if best else []
    for degrees in (90, 180, 270):
        rotated_path = image_path.with_name(f"{image_path.stem}-{degrees}.png")
        try:
            _rotated_image(image_path, rotated_path, degrees)
            text = _run_tesseract(rotated_path, languages, 6)
            if text.strip():
                rotated_candidates.append(text)
        except Exception:
            continue
    return max(rotated_candidates, key=lambda item: _text_quality(item)["score"], default="")


def extract_file_text(file_field):
    if not file_field:
        return ""
    path = Path(file_field.path)
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            return _shorten(_extract_docx(path))
        if suffix == ".pdf":
            return _shorten(_extract_pdf(path))
        if suffix in {".txt", ".md", ".csv"}:
            return _shorten(path.read_text(encoding="utf-8", errors="ignore"))
    except Exception:
        return ""
    return ""


def _source_text_hints(text):
    text = str(text or "")
    email_pattern = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
    email_candidates = [match.group(0) for match in email_pattern.finditer(text)]
    phone_match = re.search(r"\+\d{1,4}[\s().-]?\d(?:[\s().-]?\d){6,14}", text) or re.search(
        r"(?<!\d)\d(?:[\s().-]?\d){7,13}(?!\d)", text
    )
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s|,;]+", text, re.I)
    if not linkedin_match:
        linkedin_match = re.search(r"(?:linkedin|linked\s*in)[:\s/]+([A-Z0-9._-]{3,80})", text, re.I)
    github_match = re.search(r"(?:https?://)?(?:www\.)?github\.com/[A-Z0-9._-]+", text, re.I)
    portfolio_match = re.search(
        r"portfolio[^\n\r]{0,40}((?:https?://|www\.)[A-Z0-9.-]+\.[A-Z]{2,}(?:/[^\s|,;]*)?|\b[A-Z0-9-]+\.(?:com|ci|net|org|io|dev|app|co)(?:/[^\s|,;]*)?)",
        text,
        re.I,
    )
    lines = [line.strip(" |:-") for line in text.splitlines() if line.strip(" |:-")]
    handles = [match.group(1) for match in re.finditer(r"github\.com/([A-Z0-9._-]+)", text, re.I)]

    def best_email():
        if not email_candidates:
            return ""
        domain_fixes = {
            "gmall.com": "gmail.com",
            "gmai.com": "gmail.com",
            "gmail.con": "gmail.com",
        }
        common_domains = {"gmail.com", "outlook.com", "hotmail.com", "yahoo.com", "icloud.com"}
        ranked = []
        for raw in email_candidates:
            local, domain = raw.split("@", 1)
            domain = domain_fixes.get(domain.lower(), domain.lower())
            local_value = local.strip(".-_").lower()
            if local_value.startswith("+") or re.match(r"^\d{6,}", local_value):
                continue
            tld = domain.rsplit(".", 1)[-1] if "." in domain else ""
            if len(tld) > 4 or len(local_value) > 32:
                continue
            for handle in handles:
                handle_value = handle.lower()
                if local_value.startswith(handle_value[1:]) and not local_value.startswith(handle_value):
                    local_value = handle_value[:1] + local_value
                else:
                    best_cut = len(handle_value)
                    best_ratio = 0
                    min_cut = max(4, len(handle_value) - 2)
                    max_cut = min(len(local_value), len(handle_value) + 2)
                    for cut in range(min_cut, max_cut + 1):
                        ratio = SequenceMatcher(None, local_value[:cut], handle_value).ratio()
                        if ratio > best_ratio:
                            best_ratio = ratio
                            best_cut = cut
                    if best_ratio >= 0.58:
                        local_value = handle_value + local_value[best_cut:]
            score = len(local_value)
            if domain in common_domains:
                score += 20
            if raw[:1].islower():
                score += 3
            if any(char.isdigit() for char in local_value):
                score -= 2
            ranked.append((score, f"{local_value}@{domain}"))
        ranked.sort(reverse=True)
        return ranked[0][1] if ranked else ""

    name = ""
    ignored = {"cv", "curriculum vitae", "resume", "profil", "contact", "formation"}
    for line in lines[:12]:
        folded = line.lower()
        words = [word for word in re.split(r"\s+", line) if word]
        if folded in ignored or "@" in line or any(char.isdigit() for char in line):
            continue
        if 2 <= len(words) <= 5:
            name = line
            break
    possible_github = github_match.group(0) if github_match else ""
    possible_portfolio = portfolio_match.group(1) if portfolio_match and "@" not in portfolio_match.group(1) else ""
    if possible_github and possible_portfolio.lower().rstrip("/") == possible_github.lower().rstrip("/"):
        possible_portfolio = ""

    return {
        "possible_full_name": name,
        "possible_email": best_email(),
        "possible_phone": phone_match.group(0).strip() if phone_match else "",
        "possible_linkedin": linkedin_match.group(0) if linkedin_match else "",
        "possible_github": possible_github,
        "possible_portfolio": possible_portfolio,
        "line_count": len(lines),
    }


def _split_source_name(full_name):
    parts = [part for part in re.split(r"\s+", _safe_text(full_name)) if part]
    if not parts:
        return "", ""
    if len(parts) == 1:
        return "", parts[0]
    return " ".join(parts[1:]), parts[0]


def _fetch_url_text(url):
    if not url:
        return ""
    if not url.startswith(("http://", "https://")):
        return ""
    try:
        request = urllib.request.Request(url, headers={"User-Agent": "CVBuilder/1.0"})
        with urllib.request.urlopen(request, timeout=12) as response:
            content_type = response.headers.get("content-type", "")
            raw = response.read(500_000)
        text = raw.decode("utf-8", errors="ignore")
        if "html" in content_type:
            parser = _HTMLTextExtractor()
            parser.feed(text)
            return _shorten(parser.text(), 6000)
        return _shorten(text, 6000)
    except Exception:
        return ""


def _image_content(file_field):
    if not file_field:
        return None
    path = Path(file_field.path)
    mime, _ = mimetypes.guess_type(path.name)
    if not (mime or "").startswith("image/"):
        return None
    return _image_path_content(path)


def _image_path_content(path, max_size=(1600, 2200)):
    try:
        image = Image.open(path)
    except (OSError, UnidentifiedImageError):
        return None
    image = ImageOps.exif_transpose(image.convert("RGB"))
    image.thumbnail(max_size, Image.Resampling.LANCZOS)
    buffer = io.BytesIO()
    image.save(buffer, "JPEG", quality=88, optimize=True)
    data = base64.b64encode(buffer.getvalue()).decode("ascii")
    return {"type": "input_image", "image_url": f"data:image/jpeg;base64,{data}"}


def _pdf_pages_image_content(file_field, max_pages=2):
    if not file_field:
        return []
    path = Path(file_field.path)
    if path.suffix.lower() != ".pdf":
        return []
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            prefix = Path(tmpdir) / "source-page"
            completed = subprocess.run(
                ["pdftoppm", "-png", "-r", "144", "-f", "1", "-l", str(max_pages), str(path), str(prefix)],
                capture_output=True,
                text=True,
                timeout=30,
                check=False,
            )
            if completed.returncode != 0:
                return []
            contents = []
            for image_path in sorted(Path(tmpdir).glob("source-page-*.png")):
                content = _image_path_content(image_path)
                if content:
                    contents.append(content)
            return contents
    except Exception:
        return []


def _save_profile_photo(image, cv, request=None):
    image = ImageOps.exif_transpose(image.convert("RGB"))
    image.thumbnail((900, 1100), Image.Resampling.LANCZOS)
    photos_dir = Path(settings.MEDIA_ROOT) / "cvs" / "photos"
    photos_dir.mkdir(parents=True, exist_ok=True)
    photo_path = photos_dir / f"cv-{cv.id}-source-photo.png"
    image.save(photo_path, "PNG", optimize=True)
    relative = photo_path.relative_to(settings.MEDIA_ROOT).as_posix()
    media_url = settings.MEDIA_URL if settings.MEDIA_URL.startswith("/") else f"/{settings.MEDIA_URL}"
    url = f"{media_url}{relative}"
    return request.build_absolute_uri(url) if request else url


def _image_from_source_file(path):
    mime, _ = mimetypes.guess_type(path.name)
    if not (mime or "").startswith("image/"):
        return None
    try:
        return Image.open(path)
    except (OSError, UnidentifiedImageError):
        return None


def _photo_from_docx(path):
    best = None
    best_area = 0
    try:
        with ZipFile(path) as archive:
            for name in archive.namelist():
                if not name.startswith("word/media/"):
                    continue
                raw = archive.read(name)
                try:
                    image = Image.open(io.BytesIO(raw)).convert("RGB")
                except (OSError, UnidentifiedImageError):
                    continue
                width, height = image.size
                area = width * height
                if width < 100 or height < 100 or area <= best_area:
                    continue
                best = image.copy()
                best_area = area
    except Exception:
        return None
    return best


def _largest_image(paths):
    best = None
    best_area = 0
    for path in paths:
        try:
            image = Image.open(path).convert("RGB")
        except (OSError, UnidentifiedImageError):
            continue
        width, height = image.size
        area = width * height
        if width < 100 or height < 100 or area <= best_area:
            continue
        best = image.copy()
        best_area = area
    return best


def _looks_like_page_image(image):
    width, height = image.size
    ratio = height / max(width, 1)
    return width >= 900 and height >= 1200 and 1.2 <= ratio <= 1.7


def _portrait_crop_from_page(image):
    width, height = image.size
    box = (int(width * 0.035), int(height * 0.015), int(width * 0.34), int(height * 0.19))
    crop = image.crop(box).convert("RGB")
    return crop


def _photo_from_pdf(path):
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            prefix = Path(tmpdir) / "source-photo"
            completed = subprocess.run(
                ["pdfimages", "-png", "-f", "1", "-l", "1", str(path), str(prefix)],
                capture_output=True,
                text=True,
                timeout=20,
                check=False,
            )
            if completed.returncode != 0:
                return None
            image = _largest_image(Path(tmpdir).glob("source-photo-*.png"))
            if image is not None and _looks_like_page_image(image):
                return _portrait_crop_from_page(image)
            return image
    except Exception:
        return None


def _photo_file_looks_like_document(file_field):
    if not file_field:
        return False
    name = Path(file_field.name).name.lower()
    return any(token in name for token in ["piece", "pièce", "cni", "idcard", "identity", "carte", "recto", "verso"])


def _photo_from_file_field(file_field):
    if not file_field:
        return ""
    path = Path(file_field.path)
    if not path.exists():
        return None
    image = _image_from_source_file(path)
    if image is None and path.suffix.lower() == ".docx":
        image = _photo_from_docx(path)
    if image is None and path.suffix.lower() == ".pdf":
        image = _photo_from_pdf(path)
    return image


def extract_source_photo_url(cv, request=None, prefer_photo_file=False):
    image = None
    if prefer_photo_file and not _photo_file_looks_like_document(cv.photo_file):
        image = _photo_from_file_field(cv.photo_file)
    if image is None:
        image = _photo_from_file_field(cv.source_file)
    if image is None and not prefer_photo_file:
        image = _photo_from_file_field(cv.photo_file)
    if image is None:
        return ""
    return _save_profile_photo(image, cv, request=request)


def _clean_ai_data(payload):
    data = payload.get("data") or payload
    defaults = {
        "first_name": "",
        "last_name": "",
        "job_title": "",
        "photo_url": "",
        "phone": "",
        "email": "",
        "address": "",
        "linkedin": "",
        "github": "",
        "portfolio": "",
        "driving_license": "",
        "profile": "",
        "experiences": [],
        "education": [],
        "skills": [],
        "languages": [],
        "hobbies": [],
        "extra_sections": [],
    }
    cleaned = {**defaults, **{key: value for key, value in data.items() if key in defaults}}
    for key in ["experiences", "education", "skills", "languages", "hobbies", "extra_sections"]:
        if not isinstance(cleaned[key], list):
            cleaned[key] = []
    return cleaned


def _fallback_result(cv, instruction="", provider_label="IA", key_name="API_KEY", has_source=False):
    data = cv.data or {}
    merged = _clean_ai_data({"data": data})
    if not merged["profile"]:
        role = merged.get("job_title") or "profil professionnel"
        merged["profile"] = f"{role} motivé, orienté résultats, avec une présentation claire des compétences et expériences utiles au poste visé."
    return {
        "data": merged,
        "fit_summary": f"Mode hors ligne: renseigne {key_name} pour obtenir une optimisation {provider_label} complète.",
        "missing_info_questions": (
            ["Souhaitez-vous ajouter une section supplémentaire comme Projets, Certifications ou Références ?"]
            if has_source
            else [
                "Quelle est l'offre exacte ou le poste ciblé ?",
                "Quelles réalisations chiffrées faut-il mettre en avant ?",
            ]
        ),
        "change_log": ["Structure normalisée à partir des informations disponibles."],
        "template_recommendation": {
            "mode": cv.template_mode,
            "notes": "Le rendu PDF utilisera le modèle choisi dans l'application.",
        },
    }


def _schema():
    text = {"type": "string"}
    list_text = {"type": "array", "items": text}
    return {
        "type": "object",
        "additionalProperties": False,
        "required": ["data", "fit_summary", "missing_info_questions", "change_log", "template_recommendation"],
        "properties": {
            "data": {
                "type": "object",
                "additionalProperties": False,
                "required": [
                    "first_name", "last_name", "job_title", "photo_url", "phone", "email",
                    "address", "linkedin", "github", "portfolio", "driving_license", "profile", "experiences",
                    "education", "skills", "languages", "hobbies", "extra_sections"
                ],
                "properties": {
                    "first_name": text,
                    "last_name": text,
                    "job_title": text,
                    "photo_url": text,
                    "phone": text,
                    "email": text,
                    "address": text,
                    "linkedin": text,
                    "github": text,
                    "portfolio": text,
                    "driving_license": text,
                    "profile": text,
                    "experiences": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["job_title", "company", "location", "period", "type", "missions"],
                            "properties": {
                                "job_title": text,
                                "company": text,
                                "location": text,
                                "period": text,
                                "type": text,
                                "missions": list_text,
                            },
                        },
                    },
                    "education": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["degree", "institution", "location", "period"],
                            "properties": {
                                "degree": text,
                                "institution": text,
                                "location": text,
                                "period": text,
                            },
                        },
                    },
                    "skills": list_text,
                    "languages": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["language", "level"],
                            "properties": {"language": text, "level": text},
                        },
                    },
                    "hobbies": list_text,
                    "extra_sections": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["title", "items"],
                            "properties": {
                                "title": text,
                                "items": list_text,
                            },
                        },
                    },
                },
            },
            "fit_summary": text,
            "missing_info_questions": list_text,
            "change_log": list_text,
            "template_recommendation": {
                "type": "object",
                "additionalProperties": False,
                "required": ["mode", "notes"],
                "properties": {"mode": text, "notes": text},
            },
        },
    }


def _parse_response(response, provider_label="IA"):
    if isinstance(response, dict):
        choices = response.get("choices") or []
        if choices:
            message = (choices[0] or {}).get("message") or {}
            output_text = message.get("content") or ""
            if not output_text:
                raise ValueError(f"Réponse {provider_label} vide ou illisible.")
            return json.loads(output_text)

        output_text = response.get("output_text") or ""
        if not output_text:
            for item in response.get("output") or []:
                for content in item.get("content") or []:
                    if isinstance(content, dict):
                        output_text += content.get("text") or ""
        if not output_text:
            raise ValueError(f"Réponse {provider_label} vide ou illisible.")
        return json.loads(output_text)

    output_text = getattr(response, "output_text", "")
    if not output_text:
        for item in getattr(response, "output", []) or []:
            for content in getattr(item, "content", []) or []:
                value = getattr(content, "text", None)
                if value:
                    output_text += value
    if not output_text:
        raise ValueError(f"Réponse {provider_label} vide ou illisible.")
    return json.loads(output_text)


def _read_ai_error(detail, status_code, provider_label, key_name):
    code = f"{provider_label.lower()}_error"
    message = "Le service IA est momentanément indisponible. Réessaie dans quelques instants."
    try:
        payload = json.loads(detail)
        error = payload.get("error") if isinstance(payload, dict) else {}
        if isinstance(error, dict):
            code = error.get("code") or error.get("type") or code
    except json.JSONDecodeError:
        error = {}

    if code == "insufficient_quota":
        message = (
            f"Le service IA est indisponible car le quota {provider_label} du serveur est épuisé. "
            f"Ajoute du crédit ou active la facturation du compte relié à {key_name}, puis réessaie."
        )
        return message, code, 503
    if code in {"invalid_api_key", "authentication_error"} or status_code == 401:
        message = f"La clé {provider_label} configurée sur le serveur est invalide ou refusée."
        return message, code, 503
    if code == "rate_limit_exceeded" or status_code == 413:
        message = (
            "Trop de demandes IA dans la dernière minute (quota du palier gratuit). "
            "Patiente une minute puis réessaie, ou réduis la taille du CV importé."
        )
        return message, code, 429
    if status_code == 403:
        message = (
            f"{provider_label} refuse la requête depuis ce serveur pour le moment. "
            "La clé est configurée, mais l'appel a été bloqué avant une réponse IA."
        )
        return message, code, 503
    if status_code == 429:
        message = f"{provider_label} reçoit trop de demandes pour le moment. Réessaie dans quelques instants."
        return message, code, 429
    if status_code >= 500:
        message = f"{provider_label} est temporairement indisponible. Réessaie dans quelques instants."
        return message, code, 503
    return message, code, 500


def _provider_config_for(provider):
    if provider == "groq":
        return {
            "provider": provider,
            "label": "Groq",
            "endpoint": "https://api.groq.com/openai/v1/chat/completions",
            "api_key": settings.GROQ_API_KEY,
            "key_name": "GROQ_API_KEY",
            "model": settings.GROQ_MODEL,
        }
    if provider == "openai":
        return {
            "provider": provider,
            "label": "OpenAI",
            "endpoint": "https://api.openai.com/v1/responses",
            "api_key": settings.OPENAI_API_KEY,
            "key_name": "OPENAI_API_KEY",
            "model": settings.OPENAI_MODEL,
        }
    raise AIServiceError(
        f"Provider IA inconnu: {provider}. Utilise AI_PROVIDER=groq ou AI_PROVIDER=openai.",
        code="unknown_ai_provider",
        status_code=500,
    )


def _has_configured_key(config):
    key = config["api_key"]
    if not key:
        return False
    placeholder_prefixes = {
        "GROQ_API_KEY": ("gsk_xxx",),
        "OPENAI_API_KEY": ("sk-xxx",),
    }.get(config["key_name"], ())
    return not any(key.startswith(prefix) for prefix in placeholder_prefixes)


def _provider_config(preferred=None):
    return _provider_config_for(preferred or settings.AI_PROVIDER)


# Statuts HTTP transitoires que l'on peut réessayer (Cloudflare 1010=403, rate-limit, pannes).
_RETRYABLE_STATUSES = {403, 408, 425, 429, 500, 502, 503, 504}
_AI_MAX_ATTEMPTS = 3


def _ai_responses_create(payload, config):
    body = json.dumps(payload).encode("utf-8")
    request = urllib.request.Request(
        config["endpoint"],
        data=body,
        method="POST",
        headers={
            "Authorization": f"Bearer {config['api_key']}",
            "Accept": "application/json",
            "Content-Type": "application/json",
            # User-Agent type navigateur : évite les blocages Cloudflare (erreur 1010)
            # qui rejettent les clients « Python-urllib » / non navigateurs.
            "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36",
        },
    )

    last_error = None
    for attempt in range(_AI_MAX_ATTEMPTS):
        try:
            with urllib.request.urlopen(request, timeout=settings.OPENAI_REQUEST_TIMEOUT) as response:
                return json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="ignore")
            message, code, response_status = _read_ai_error(detail, exc.code, config["label"], config["key_name"])
            last_error = AIServiceError(message, code=code, status_code=response_status, raw_detail=detail)
            # Une clé invalide / quota épuisé ne se résout pas en réessayant : on arrête.
            if exc.code not in _RETRYABLE_STATUSES or code in {"insufficient_quota", "invalid_api_key", "authentication_error"}:
                raise last_error from exc
        except (urllib.error.URLError, TimeoutError, OSError) as exc:
            last_error = AIServiceError(
                f"Impossible de contacter {config['label']} depuis le serveur. Vérifie la connexion réseau puis réessaie.",
                code=f"{config['provider']}_network_error",
                status_code=503,
                raw_detail=str(exc),
            )

        if attempt < _AI_MAX_ATTEMPTS - 1:
            time.sleep(1.5 * (attempt + 1))  # backoff: 1.5s, 3s

    raise last_error


def rewrite_cv_text(text, kind="texte", max_words=55):
    """Propose une version plus courte, fluide et fidèle d'un texte de CV.
    Ton humain, jamais marqué « IA ». Renvoie uniquement le texte réécrit."""
    text = _safe_text(text)
    if not text:
        raise AIServiceError("Aucun texte à réécrire.", code="empty_text", status_code=400)

    config = _provider_config()
    if not _has_configured_key(config):
        raise AIServiceError(
            f"La clé {config['label']} n'est pas configurée sur le serveur.",
            code="missing_ai_key",
            status_code=503,
        )

    labels = {
        "profile": "accroche / phrase de profil de CV",
        "mission": "description d'une mission ou d'une réalisation",
        "experience": "description d'expérience professionnelle",
    }
    label = labels.get(kind, "texte de CV")
    system_prompt = (
        f"Tu es un coach emploi francophone. Réécris la {label} fournie de façon plus courte, "
        "fluide, claire et naturelle, avec un ton humain et professionnel (garde la même personne "
        "grammaticale que l'original). Reste strictement fidèle aux faits : n'invente rien, ne "
        "supprime aucune information importante, ne change pas les chiffres, dates ou noms. "
        f"Vise environ {int(max_words)} mots maximum. "
        "Ne mentionne jamais que tu es une IA ni que c'est une réécriture. "
        "Pas de guillemets, pas de préambule, pas de liste à puces : réponds uniquement avec le texte final."
    )

    if config["provider"] == "groq":
        payload = {
            "model": config["model"],
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text},
            ],
            "temperature": 0.4,
            "max_completion_tokens": 700,
            "reasoning_effort": "low",
        }
        response = _ai_responses_create(payload, config)
        try:
            content = response["choices"][0]["message"]["content"]
        except (KeyError, IndexError, TypeError):
            raise AIServiceError("Réponse IA inattendue.", code="bad_ai_response", status_code=502)
    else:
        payload = {
            "model": config["model"],
            "input": [
                {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
                {"role": "user", "content": [{"type": "input_text", "text": text}]},
            ],
        }
        response = _ai_responses_create(payload, config)
        content = response.get("output_text") or _parse_response(response, provider_label=config["label"]).get("raw", "")

    return _safe_text(content).strip().strip('"').strip()


def write_profile(data, job_offer=""):
    """Rédige (ou améliore) une accroche de profil de 2-3 lignes à partir des
    informations du CV. Si une offre est fournie, oriente le profil vers elle
    (sans inventer). Fidèle, brève, humaine, sans mention d'IA."""
    data = data or {}
    job_offer = _shorten(job_offer, 1800)
    config = _provider_config()
    if not _has_configured_key(config):
        raise AIServiceError(
            f"La clé {config['label']} n'est pas configurée sur le serveur.",
            code="missing_ai_key",
            status_code=503,
        )

    experiences = []
    for exp in _clean_items(data.get("experiences"))[:4]:
        experiences.append({
            "poste": _safe_text(exp.get("job_title")),
            "entreprise": _safe_text(exp.get("company")),
            "missions": [_safe_text(m) for m in (exp.get("missions") or []) if _safe_text(m)][:3],
        })
    summary = {
        "intitule": _safe_text(data.get("job_title")),
        "profil_actuel": _safe_text(data.get("profile")),
        "experiences": experiences,
        "formations": [
            {"diplome": _safe_text(e.get("degree")), "etablissement": _safe_text(e.get("institution"))}
            for e in _clean_items(data.get("education"))[:4]
        ],
        "competences": [_safe_text(s) for s in _clean_items(data.get("skills")) if _safe_text(s)][:12],
    }
    if not (summary["intitule"] or summary["experiences"] or summary["formations"] or summary["competences"]):
        raise AIServiceError(
            "Renseigne d'abord ton intitulé, tes expériences ou tes diplômes : l'IA s'en sert pour rédiger le profil.",
            code="not_enough_data",
            status_code=400,
        )

    if job_offer:
        summary["offre_ciblee"] = job_offer
    system_prompt = (
        "Tu es un coach emploi francophone. Rédige une accroche de profil de CV COURTE : 2 à 3 lignes, "
        "35 à 50 mots maximum, à la première personne, professionnelle, fluide et humaine, à partir "
        "UNIQUEMENT des informations fournies. N'invente aucun diplôme, chiffre, entreprise ni compétence. "
        "Si une offre ciblée est fournie, oriente l'accroche vers ce poste en mettant en avant les "
        "informations pertinentes du candidat, sans rien inventer. "
        "Si un profil actuel est fourni, améliore-le sans en changer le sens. "
        "Ne mentionne jamais que tu es une IA. Pas de guillemets, pas de titre, pas de préambule : "
        "réponds uniquement avec le texte du profil."
    )
    user_message = json.dumps(summary, ensure_ascii=False)
    if config["provider"] == "groq":
        payload = {
            "model": config["model"],
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            "temperature": 0.5,
            "max_completion_tokens": 700,
            "reasoning_effort": "low",
        }
        response = _ai_responses_create(payload, config)
        try:
            content = response["choices"][0]["message"]["content"]
        except (KeyError, IndexError, TypeError):
            raise AIServiceError("Réponse IA inattendue.", code="bad_ai_response", status_code=502)
    else:
        payload = {
            "model": config["model"],
            "input": [
                {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
                {"role": "user", "content": [{"type": "input_text", "text": user_message}]},
            ],
        }
        response = _ai_responses_create(payload, config)
        content = response.get("output_text") or ""

    return _safe_text(content).strip().strip('"').strip()


def correct_cv_data(data):
    """Correction globale du CV : orthographe, accents, ponctuation, majuscules et
    format des périodes/dates (ex: « 2025 À 2026 » -> « 2025 - 2026 »). Ne change
    pas le sens, ne supprime ni n'invente rien. Renvoie les données corrigées."""
    data = data or {}
    config = _provider_config()
    if not _has_configured_key(config):
        raise AIServiceError(
            f"La clé {config['label']} n'est pas configurée sur le serveur.",
            code="missing_ai_key",
            status_code=503,
        )

    system_prompt = (
        "Tu es un correcteur de CV francophone. Tu reçois les données d'un CV en JSON. "
        "Corrige UNIQUEMENT : l'orthographe, les accents (é, è, à, ç, ê…), la ponctuation, les majuscules en début "
        "de phrase, et le FORMAT des périodes/dates (ex: « 2025 À 2026 », « 2025 a 2026 » -> « 2025 - 2026 » ; "
        "« janvier 2024 a janvier 2026 » -> « Janvier 2024 - Janvier 2026 »). "
        "Tu ne dois RIEN inventer, RIEN supprimer, RIEN raccourcir, ni changer le sens. "
        "Conserve TOUTES les expériences, missions, formations, diplômes, compétences, langues et sections. "
        "Renvoie EXACTEMENT les mêmes données, corrigées, dans le schéma JSON demandé."
    )
    user_message = json.dumps({"cv_data": data}, ensure_ascii=False)
    if config["provider"] == "groq":
        estimated_prompt_tokens = (len(system_prompt) + len(user_message)) // 4 + 400
        max_out = max(1800, min(6000, settings.GROQ_TPM_LIMIT - estimated_prompt_tokens))
        payload = {
            "model": config["model"],
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            "response_format": {"type": "json_schema", "json_schema": {"name": "cv_ai_result", "schema": _schema()}},
            "temperature": 0.1,
            "max_completion_tokens": max_out,
            "reasoning_effort": "low",
        }
    else:
        content = [
            {"type": "input_text", "text": system_prompt},
            {"type": "input_text", "text": user_message},
        ]
        payload = {
            "model": config["model"],
            "input": [{"role": "user", "content": content}],
            "text": {"format": {"type": "json_schema", "name": "cv_ai_result", "strict": True, "schema": _schema()}},
        }
    response = _ai_responses_create(payload, config)
    corrected = _clean_ai_data(_parse_response(response, provider_label=config["label"]))
    # On préserve les champs gérés côté front (jamais touchés par l'IA).
    for key in ("photo_url", "enabled_sections", "section_order"):
        if data.get(key) and not corrected.get(key):
            corrected[key] = data[key]
    corrected["photo_url"] = data.get("photo_url") or corrected.get("photo_url", "")
    return corrected


def improve_cv(cv, instruction=""):
    config = _provider_config()
    if not _has_configured_key(config):
        return _fallback_result(
            cv,
            instruction,
            provider_label=config["label"],
            key_name=config["key_name"],
            has_source=bool(cv.source_file),
        )

    source_text = extract_file_text(cv.source_file)
    source_quality = _text_quality(source_text)
    if cv.source_file and not source_quality["usable"] and config["provider"] != "openai":
        raise AIServiceError(
            "PDF illisible: le fichier ne contient pas assez de texte exploitable. "
            "Importe un PDF texte ou installe l'OCR local Tesseract pour lire les scans.",
            code="source_cv_unreadable",
            status_code=400,
        )
    source_hints = _source_text_hints(source_text)
    offer_file_text = extract_file_text(cv.job_offer_file)
    offer_url_text = _fetch_url_text(cv.job_offer_url)
    image = _image_content(cv.job_offer_file) if config["provider"] == "openai" else None
    source_images = (
        _pdf_pages_image_content(cv.source_file)
        if config["provider"] == "openai" and cv.source_file and not source_text
        else []
    )

    system_prompt = (
        "Tu es un expert RH francophone spécialisé dans les CV propres et ATS-friendly. "
        "RÈGLE ABSOLUE DE FIDÉLITÉ: tu ne dois OMETTRE AUCUNE information présente dans uploaded_cv_text. "
        "Reprends TOUTES les expériences, TOUTES les missions/tâches de chaque expérience, TOUTES les formations, TOUS les diplômes, TOUTES les compétences, TOUTES les langues et TOUS les projets, sans en sauter un seul. "
        "Si une expérience contient 6 missions, tu restitues les 6. Tu peux corriger l'orthographe et reformuler légèrement pour la clarté, mais tu ne dois jamais supprimer, fusionner abusivement ni raccourcir au point de perdre un mot d'information. "
        "Ne perds aucun chiffre, date, nom d'entreprise, intitulé de poste, établissement ni résultat. "
        "Tu remplis exclusivement des champs JSON structurés; tu ne dois jamais coller le texte brut de l'ancien CV dans un seul champ. "
        "Tu extrais chaque information dans le champ correspondant: identité, contact (téléphone, email, ville, LinkedIn, GitHub, portfolio), profil, expériences, formations, compétences, langues, loisirs et sections supplémentaires. "
        "Quand uploaded_cv_text contient des retours ligne, utilise-les comme indices de rubriques et de hiérarchie. "
        "Une expérience doit devenir un objet experiences avec poste, entreprise, période, lieu et missions; une formation doit devenir un objet education. "
        "Les listes de compétences doivent être séparées en items courts, pas laissées dans un paragraphe. "
        "Le rendu visuel final sera imposé par le modèle choisi côté application. "
        "Tu dois respecter strictement le modèle sélectionné et ne jamais proposer un autre modèle. "
        "L'ancien CV, s'il existe, sert uniquement de source de données et de photo, jamais de modèle graphique. "
        "Tu dois optimiser la formulation du CV pour l'offre cible sans inventer de diplômes, entreprises, dates, chiffres ou expériences, et sans rien supprimer de ce qui existe déjà. "
        "Pour les CV francophones où le nom complet est écrit en ligne, considère que le premier mot est le nom de famille et que les mots suivants sont les prénoms; par exemple ASSANVO BROU ANTOINE donne last_name=ASSANVO et first_name=BROU ANTOINE. "
        "Si uploaded_cv_text est fourni, extrais les informations disponibles et ne pose pas de questions générales; "
        "ne redemande jamais nom, contact, formation, expériences ou compétences si ces informations apparaissent dans uploaded_cv_text; "
        "pose uniquement une question courte sur les sections supplémentaires à ajouter si nécessaire. "
        "Si aucun ancien CV n'est fourni, pose les questions nécessaires dans missing_info_questions. "
        "Réponds uniquement avec le JSON demandé."
    )
    user_payload = {
        "current_cv_data": cv.data or {},
        "uploaded_cv_text": source_text,
        "uploaded_cv_text_quality": source_quality,
        "uploaded_cv_detected_hints": source_hints,
        "name_order_rule": "first token is last_name, remaining tokens are first_name when the source CV writes the full name in one line",
        "uploaded_cv_image_pages": len(source_images),
        "job_offer_url": cv.job_offer_url,
        "job_offer_url_text": offer_url_text,
        "job_offer_text": cv.job_offer_text,
        "job_offer_file_text": offer_file_text,
        "template_mode": cv.template_mode,
        "selected_template": {
            "id": cv.template_id,
            "name": getattr(cv.template, "name", ""),
            "slug": getattr(cv.template, "slug", ""),
            "category": getattr(cv.template, "category", ""),
            "design_contract": template_design_contract(cv.template),
        },
        "user_instruction": instruction,
        "has_uploaded_cv": bool(cv.source_file),
    }
    content = [
        {"type": "input_text", "text": system_prompt},
        {"type": "input_text", "text": json.dumps(user_payload, ensure_ascii=False)},
    ]
    if image:
        content.append(image)
    content.extend(source_images)

    if not config["model"]:
        raise AIServiceError(f"{config['label']}_MODEL manquant dans backend/.env.", code="missing_ai_model")

    if config["provider"] == "groq":
        user_message = json.dumps(user_payload, ensure_ascii=False)
        # Le palier gratuit Groq plafonne à ~8000 tokens/minute (entrée + sortie).
        # On calcule dynamiquement la marge de sortie pour ne jamais dépasser cette
        # limite (sinon HTTP 413), tout en laissant assez de place pour extraire
        # toutes les sections sans troncature. ~4 caractères = 1 token.
        estimated_prompt_tokens = (len(system_prompt) + len(user_message)) // 4 + 400
        groq_max_completion = max(1800, min(6000, settings.GROQ_TPM_LIMIT - estimated_prompt_tokens))
        request_payload = {
            "model": config["model"],
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            "response_format": {
                "type": "json_schema",
                "json_schema": {
                    "name": "cv_ai_result",
                    "schema": _schema(),
                },
            },
            "temperature": 0.2,
            "max_completion_tokens": groq_max_completion,
            "reasoning_effort": "low",
        }
    else:
        request_payload = {
            "model": config["model"],
            "input": [{"role": "user", "content": content}],
            "text": {
                "format": {
                    "type": "json_schema",
                    "name": "cv_ai_result",
                    "strict": True,
                    "schema": _schema(),
                }
            },
        }

    response = _ai_responses_create(request_payload, config)
    payload = _parse_response(response, provider_label=config["label"])
    payload["data"] = _clean_ai_data(payload)
    if source_text:
        data = payload["data"]
        if not data.get("email") and source_hints.get("possible_email"):
            data["email"] = source_hints["possible_email"]
        if not data.get("phone") and source_hints.get("possible_phone"):
            data["phone"] = source_hints["possible_phone"]
        if not data.get("linkedin") and source_hints.get("possible_linkedin"):
            data["linkedin"] = source_hints["possible_linkedin"]
        if not data.get("github") and source_hints.get("possible_github"):
            data["github"] = source_hints["possible_github"]
        if not data.get("portfolio") and source_hints.get("possible_portfolio"):
            data["portfolio"] = source_hints["possible_portfolio"]
        if source_hints.get("possible_full_name") and not (data.get("first_name") or data.get("last_name")):
            first_name, last_name = _split_source_name(source_hints["possible_full_name"])
            data["first_name"] = first_name
            data["last_name"] = last_name
    if cv.source_file and not payload["data"].get("photo_url"):
        payload["data"]["photo_url"] = (cv.data or {}).get("photo_url", "")
    if cv.source_file:
        data = payload["data"]
        contact_questions = []
        if not data.get("phone"):
            contact_questions.append("Quel téléphone doit apparaître dans la section Contact ?")
        if not data.get("email"):
            contact_questions.append("Quel email doit apparaître dans la section Contact ?")
        if not (data.get("linkedin") or data.get("github") or data.get("portfolio")):
            contact_questions.append("Souhaitez-vous ajouter un lien LinkedIn, GitHub ou portfolio ?")
        questions = payload.get("missing_info_questions") or []
        section_questions = [
            question for question in questions
            if any(token in str(question).lower() for token in ["section", "ajout", "ajouter", "projet", "certification"])
        ]
        payload["missing_info_questions"] = (contact_questions + section_questions)[:3] or [
            "Souhaitez-vous ajouter une section supplémentaire comme Projets, Certifications ou Références ?"
        ]
    return payload


def merge_ai_result(cv, result, instruction=""):
    current = cv.data or {}
    ai_data = result.get("data") or {}
    preserved = {
        key: value
        for key, value in current.items()
        if key in {"enabled_sections", "section_order"}
    }
    merged = {**ai_data}
    for key in [
        "first_name", "last_name", "job_title", "phone", "email", "address", "linkedin",
        "github", "portfolio", "driving_license", "photo_url",
    ]:
        if not _safe_text(merged.get(key)) and _safe_text(current.get(key)):
            merged[key] = current[key]
    for key in ["experiences", "education", "skills", "languages", "hobbies", "extra_sections"]:
        if not _clean_items(merged.get(key)) and _clean_items(current.get(key)):
            merged[key] = current[key]

    cv.data = {**merged, **preserved}
    cv.ai_data = result
    cv.ai_error = ""
    cv.ai_status = cv.AI_STATUS_READY
    cv.status = cv.STATUS_AI_READY
    messages = list(cv.ai_messages or [])
    messages.append({
        "instruction": instruction,
        "fit_summary": result.get("fit_summary", ""),
        "missing_info_questions": result.get("missing_info_questions", []),
        "change_log": result.get("change_log", []),
    })
    cv.ai_messages = messages[-12:]
    cv.save(update_fields=["data", "ai_data", "ai_error", "ai_status", "status", "ai_messages", "updated_at"])
    return cv
